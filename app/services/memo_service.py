import io
import uuid
from datetime import datetime, timezone

from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select

from app.core.database import supabase
from app.models.finding import Finding
from app.models.memo import AuditMemo
from app.models.opd import OpdUnit


def _add_heading(doc: DocxDocument, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _add_table_row(table, cells: list[str]) -> None:
    row = table.add_row()
    for i, cell_text in enumerate(cells):
        row.cells[i].text = cell_text


def build_docx(findings: list[Finding], opd: OpdUnit | None, fiscal_year: int, memo_number: str) -> bytes:
    doc = DocxDocument()

    # Header
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("MEMO AUDIT KEUANGAN DAERAH")
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph(f"Nomor Memo   : {memo_number}")
    doc.add_paragraph(f"Tanggal       : {datetime.now(timezone.utc).strftime('%d %B %Y')}")
    doc.add_paragraph(f"Tahun Anggaran: {fiscal_year}")
    if opd:
        doc.add_paragraph(f"Instansi      : {opd.name} — {opd.kabupaten}")
    doc.add_paragraph()

    # Section 1: Dasar Pemeriksaan
    _add_heading(doc, "I. DASAR PEMERIKSAAN", level=2)
    doc.add_paragraph(
        "Pemeriksaan ini dilaksanakan berdasarkan:\n"
        "1. Perpres 12/2021 tentang Pengadaan Barang/Jasa Pemerintah\n"
        "2. Permendagri 77/2020 tentang Pengelolaan Keuangan Daerah\n"
        "3. PP 12/2019 tentang Pengelolaan Keuangan Daerah"
    )

    # Section 2: Ringkasan Eksekutif
    _add_heading(doc, "II. RINGKASAN EKSEKUTIF", level=2)
    red = [f for f in findings if f.risk_level == "red"]
    yellow = [f for f in findings if f.risk_level == "yellow"]
    compliance = [f for f in findings if f.source == "compliance_scan"]
    anomaly = [f for f in findings if f.source == "procurement_anomaly"]

    doc.add_paragraph(
        f"Total temuan: {len(findings)}\n"
        f"  • Risiko Tinggi (Merah) : {len(red)}\n"
        f"  • Risiko Sedang (Kuning): {len(yellow)}\n"
        f"  • Temuan Kepatuhan Regulasi: {len(compliance)}\n"
        f"  • Temuan Anomali Pengadaan : {len(anomaly)}"
    )

    # Section 3: Temuan Kepatuhan Regulasi
    _add_heading(doc, "III. TEMUAN KEPATUHAN REGULASI", level=2)
    if compliance:
        table = doc.add_table(rows=1, cols=5)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        for i, h in enumerate(["No", "Uraian Temuan", "Referensi Regulasi", "Nilai", "Rekomendasi"]):
            hdr[i].text = h
        for idx, f in enumerate(compliance, start=1):
            refs = ""
            if f.regulation_refs and isinstance(f.regulation_refs, list):
                refs = "; ".join(
                    f"{r.get('peraturan', '')} {r.get('pasal', '')}"
                    for r in f.regulation_refs
                )
            _add_table_row(table, [
                str(idx),
                f.title,
                refs,
                "-",
                f.description[:200] if f.description else "-",
            ])
    else:
        doc.add_paragraph("Tidak ada temuan kepatuhan regulasi.")

    # Section 4: Temuan Anomali Pengadaan
    _add_heading(doc, "IV. TEMUAN ANOMALI PENGADAAN", level=2)
    if anomaly:
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        for i, h in enumerate(["No", "Jenis Anomali", "Deskripsi", "Penjelasan AI"]):
            hdr[i].text = h
        for idx, f in enumerate(anomaly, start=1):
            _add_table_row(table, [
                str(idx),
                f.finding_type,
                f.description[:200] if f.description else "-",
                (f.ai_explanation or "-")[:200],
            ])
    else:
        doc.add_paragraph("Tidak ada temuan anomali pengadaan.")

    # Section 5: Rekomendasi
    _add_heading(doc, "V. REKOMENDASI TINDAK LANJUT", level=2)
    if red:
        doc.add_paragraph(
            f"Terdapat {len(red)} temuan risiko tinggi yang memerlukan tindak lanjut segera:\n"
            + "\n".join(f"  {i+1}. {f.title}" for i, f in enumerate(red[:10]))
        )
    else:
        doc.add_paragraph("Tidak ada temuan risiko tinggi yang memerlukan tindak lanjut segera.")

    # Section 6: Penutup
    _add_heading(doc, "VI. PENUTUP", level=2)
    doc.add_paragraph(
        "Memo audit ini disusun secara otomatis oleh sistem PANTAU (Platform AI Transparansi "
        "dan Audit Keuangan Daerah) berdasarkan analisis dokumen keuangan dan pengadaan. "
        "Temuan dalam memo ini bersifat indikatif dan perlu ditindaklanjuti dengan verifikasi lapangan."
    )

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


async def generate_memo(
    finding_ids: list[uuid.UUID],
    fiscal_year: int,
    opd_id: uuid.UUID | None,
    fmt: str,
    db: AsyncSession,
) -> AuditMemo:
    findings_result = await db.execute(
        select(Finding).where(Finding.id.in_(finding_ids))
    )
    findings = findings_result.scalars().all()

    opd = None
    if opd_id:
        opd_result = await db.execute(select(OpdUnit).where(OpdUnit.id == opd_id))
        opd = opd_result.scalar_one_or_none()

    memo_number = f"PANTAU/{fiscal_year}/{datetime.now(timezone.utc).strftime('%m%d%H%M')}"

    docx_bytes = build_docx(findings, opd, fiscal_year, memo_number)

    memo = AuditMemo(
        opd_id=opd_id,
        fiscal_year=fiscal_year,
        finding_ids=[str(fid) for fid in finding_ids],
        format=fmt,
        memo_number=memo_number,
    )
    db.add(memo)
    await db.flush()

    storage_path = f"{memo.id}/memo_{fiscal_year}.docx"
    supabase.storage.from_("memos").upload(
        path=storage_path,
        file=docx_bytes,
        file_options={"content-type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"},
    )
    memo.storage_path = storage_path
    await db.commit()
    await db.refresh(memo)
    return memo
